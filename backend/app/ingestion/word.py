"""DOCX parsing via python-docx."""

import io

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.core.config import settings
from app.core.logging import get_logger
from app.ingestion import ocr
from app.ingestion.archives import guard_ooxml
from app.ingestion.base import DocumentParser, ParsedDocument, ParsedPage, UnparsableDocumentError
from app.ingestion.tables import render_table

logger = get_logger(__name__)


class DocxParser(DocumentParser):
    content_types = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",)
    extensions = (".docx",)

    def parse(self, data: bytes, *, filename: str) -> ParsedDocument:
        guard_ooxml(data, filename=filename)

        try:
            document = docx.Document(io.BytesIO(data))
        except Exception as exc:  # noqa: BLE001 - python-docx raises several types
            raise UnparsableDocumentError(
                f"Could not read {filename} as a Word document. "
                "If it is an older .doc file, save it as .docx first."
            ) from exc

        # A Word file stores no page boundaries. Pagination is computed by the
        # renderer from font metrics, margins and paper size, so "page 7" does
        # not exist in the file at all. What *is* stored is an explicit page
        # break where the author forced one — so those are the only honest
        # boundaries available, and a document without any is one page.
        pages: list[list[str]] = [[]]
        # Image relationship ids per section, so OCR text is attributed to the
        # same page-break section as the paragraph the image sits in. A Word
        # file has no page numbers, but it does record where the author forced
        # a break, and that is the only honest location available.
        section_images: list[list[str]] = [[]]

        for block in _iter_blocks(document):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    pages[-1].append(_style_prefix(block) + text)
                section_images[-1].extend(_image_relationship_ids(block))
                if _has_page_break(block):
                    pages.append([])
                    section_images.append([])
            else:
                rendered = render_table([[cell.text for cell in row.cells] for row in block.rows])
                if rendered:
                    pages[-1].append(rendered)

        parsed_pages = [
            ParsedPage(number=index, text="\n\n".join(blocks).strip())
            for index, blocks in enumerate(pages, start=1)
            if any(block.strip() for block in blocks)
        ]

        if not parsed_pages:
            raise UnparsableDocumentError(f"No text could be extracted from {filename}.")

        # Renumber: a document whose only content follows a leading page break
        # would otherwise start at page 2.
        parsed_pages = [
            ParsedPage(number=index, text=page.text)
            for index, page in enumerate(parsed_pages, start=1)
        ]

        ocr_pages = _ocr_images(document, section_images, parsed_pages)
        if ocr_pages:
            # Appended rather than merged: the OCR text keeps its own
            # content_type, and merging would relabel the whole section.
            parsed_pages = parsed_pages + ocr_pages

        core = document.core_properties
        logger.info(
            "docx_parsed",
            extra={
                "source_name": filename,
                "pages": len(parsed_pages),
                "ocr_sections": len(ocr_pages),
            },
        )

        return ParsedDocument(
            pages=parsed_pages,
            metadata={
                "title": core.title or None,
                "author": core.author or None,
                "source_page_count": len(parsed_pages),
            },
        )


def _iter_blocks(document: DocxDocument):
    """Yield paragraphs and tables in document order.

    `document.paragraphs` and `document.tables` are separate collections, so
    reading them one after the other silently reorders the document — every
    table ends up after all the prose. Walking the body XML preserves the
    order the author wrote.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _has_page_break(paragraph: Paragraph) -> bool:
    for run in paragraph.runs:
        for break_element in run._element.findall(qn("w:br")):
            if break_element.get(qn("w:type")) == "page":
                return True
    return False


def _style_prefix(paragraph: Paragraph) -> str:
    """Mark headings with markdown hashes.

    Structure carries meaning that plain text loses. Keeping headings visible
    helps the chunker split on section boundaries and helps the model tell a
    section title from a sentence.
    """
    name = (paragraph.style.name or "").lower()
    if name.startswith("heading"):
        digits = "".join(character for character in name if character.isdigit())
        level = min(int(digits), 6) if digits else 1
        return "#" * level + " "
    if name == "title":
        return "# "
    return ""


def _image_relationship_ids(paragraph: Paragraph) -> list[str]:
    """Relationship ids of images anchored in this paragraph.

    Read from the paragraph's own XML rather than from the document's image
    parts as a whole, because the position is the point: an image found this
    way can be attributed to the section it appears in, while
    `document.part.related_parts` gives an unordered bag with no location at
    all.
    """
    blips = paragraph._element.findall(".//" + qn("a:blip"))
    return [embed for blip in blips if (embed := blip.get(qn("r:embed")))]


def _ocr_images(
    document: DocxDocument,
    section_images: list[list[str]],
    parsed_pages: list[ParsedPage],
) -> list[ParsedPage]:
    """Read text out of the images embedded in each section."""
    if not ocr.is_available() or not any(section_images):
        return []

    text_by_number = {page.number: page.text for page in parsed_pages}
    related = document.part.related_parts
    budget = settings.ocr_max_images_per_document
    results: list[ParsedPage] = []

    for index, relationship_ids in enumerate(section_images, start=1):
        recognised: list[str] = []

        for relationship_id in relationship_ids:
            if budget <= 0:
                break
            part = related.get(relationship_id)
            if part is None or not str(part.content_type).startswith("image/"):
                continue
            budget -= 1

            result = ocr.extract_text(part.blob, context=f"docx section {index}")
            if result.usable and not ocr.is_redundant(result.text, text_by_number.get(index, "")):
                recognised.append(result.text)

        if recognised:
            # Clamped to a section that survived the empty-section filter, so
            # the citation points somewhere that exists.
            fallback = max(text_by_number) if text_by_number else 1
            number = index if index in text_by_number else fallback
            results.append(
                ParsedPage(
                    number=number,
                    text="\n\n".join(recognised),
                    metadata={"content_type": ocr.OCR_CONTENT_TYPE},
                )
            )

    return results
